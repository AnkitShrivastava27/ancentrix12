import os
import logging
from datetime import datetime
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, BackgroundTasks
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.database import get_db
from app.core.security import get_current_active_user
from app.models.models import Company, KnowledgeDocument
from app.services.llm.rag_service import rag_service

router = APIRouter()
logger = logging.getLogger(__name__)

ALLOWED_TYPES = {
    "application/pdf": "pdf",
    "text/plain": "txt",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/csv": "csv",
}


async def _company(user_id: str, db: AsyncSession) -> Company:
    # v2 single-tenant: ignore user_id, just get the one company
    r = await db.execute(select(Company).limit(1))
    c = r.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Company not found. Please complete setup in Settings.")
    return c


@router.get("/")
async def list_documents(
    current_user=Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db),
):
    company = await _company(current_user.id, db)
    r = await db.execute(
        select(KnowledgeDocument)
        .where(KnowledgeDocument.company_id == company.id)
        .order_by(KnowledgeDocument.created_at.desc())
    )
    docs = r.scalars().all()
    return [
        {
            "id": d.id, "filename": d.filename, "file_type": d.file_type,
            "status": d.status, "chunks_count": d.chunks_count,
            "file_size": d.file_size, "error_msg": d.error_msg,
            "created_at": d.created_at,
        }
        for d in docs
    ]


@router.get("/debug")
async def debug_rag(
    current_user=Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Debug endpoint — call this to check what's actually stored in ChromaDB
    for your company. If total_chunks=0, the PDF wasn't ingested properly.
    Visit: GET /api/v1/knowledge/debug
    """
    company = await _company(current_user.id, db)
    import asyncio
    result = await asyncio.get_event_loop().run_in_executor(
        None, rag_service.debug_collection, company.id
    )
    # Also show DB records
    r = await db.execute(
        select(KnowledgeDocument)
        .where(KnowledgeDocument.company_id == company.id)
    )
    docs = r.scalars().all()
    result["db_documents"] = [
        {"filename": d.filename, "status": d.status,
         "chunks_count": d.chunks_count, "error_msg": d.error_msg}
        for d in docs
    ]
    return result


@router.post("/reindex/{doc_id}")
async def reindex_document(
    doc_id: str,
    background_tasks: BackgroundTasks,
    current_user=Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db),
):
    """Re-process a document that failed or has 0 chunks."""
    company = await _company(current_user.id, db)
    r = await db.execute(
        select(KnowledgeDocument).where(
            KnowledgeDocument.id == doc_id,
            KnowledgeDocument.company_id == company.id,
        )
    )
    doc = r.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found")
    if not doc.file_path or not os.path.exists(doc.file_path):
        raise HTTPException(400, "Original file not found on disk — please re-upload")

    doc.status = "pending"
    doc.error_msg = None
    await db.commit()

    background_tasks.add_task(
        _process_document, doc.id, company.id, doc.file_path, doc.file_type
    )
    return {"status": "reindexing", "doc_id": doc_id}


@router.post("/upload")
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    current_user=Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db),
):
    company = await _company(current_user.id, db)

    content_type = file.content_type or ""
    file_ext = ALLOWED_TYPES.get(content_type)
    if not file_ext:
        fname = file.filename or ""
        if fname.endswith(".pdf"):    file_ext = "pdf"
        elif fname.endswith(".txt"):  file_ext = "txt"
        elif fname.endswith(".docx"): file_ext = "docx"
        elif fname.endswith(".csv"):  file_ext = "csv"
        else:
            raise HTTPException(400, "Unsupported file type. Allowed: PDF, TXT, DOCX, CSV")

    content = await file.read()
    MAX_FILE_SIZE_MB = 20
    if len(content) > MAX_FILE_SIZE_MB * 1024 * 1024:
        raise HTTPException(400, f"File too large. Max {MAX_FILE_SIZE_MB}MB")

    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    safe_name = file.filename.replace(" ", "_").replace("/", "_")
    file_path = os.path.join(settings.UPLOAD_DIR, f"{company.id}_{safe_name}")
    with open(file_path, "wb") as f:
        f.write(content)

    doc = KnowledgeDocument(
        company_id=company.id,
        filename=file.filename,
        file_type=file_ext,
        file_path=file_path,
        file_size=len(content),
        status="pending",
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    background_tasks.add_task(_process_document, doc.id, company.id, file_path, file_ext)

    return {"id": doc.id, "filename": doc.filename, "status": "processing"}


@router.delete("/{doc_id}")
async def delete_document(
    doc_id: str,
    current_user=Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db),
):
    company = await _company(current_user.id, db)
    r = await db.execute(
        select(KnowledgeDocument).where(
            KnowledgeDocument.id == doc_id,
            KnowledgeDocument.company_id == company.id,
        )
    )
    doc = r.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found")

    await rag_service.delete_document(company.id, doc_id)

    if doc.file_path and os.path.exists(doc.file_path):
        os.remove(doc.file_path)

    await db.delete(doc)
    await db.commit()
    return {"deleted": True}


async def _process_document(doc_id: str, company_id: str, file_path: str, file_type: str):
    """Extract text from document and ingest into ChromaDB."""
    from app.core.database import AsyncSessionLocal

    async with AsyncSessionLocal() as db:
        doc = await db.get(KnowledgeDocument, doc_id)
        if not doc:
            return

        try:
            doc.status = "processing"
            await db.commit()

            text = _extract_text(file_path, file_type)
            if not text.strip():
                doc.status = "failed"
                doc.error_msg = "No text could be extracted from file"
                await db.commit()
                logger.warning(f"No text extracted from {file_path}")
                return

            logger.info(f"Extracted {len(text)} chars from {doc.filename}")

            chunks = await rag_service.ingest_text(
                company_id=company_id,
                doc_id=doc_id,
                text=text,
                metadata={"filename": doc.filename, "file_type": file_type},
            )

            doc.status = "completed"
            doc.chunks_count = chunks
            doc.updated_at = datetime.utcnow()
            await db.commit()
            logger.info(f"Ingested {doc.filename}: {chunks} chunks into ChromaDB")

        except Exception as e:
            logger.error(f"Document processing error for {doc.filename}: {e}", exc_info=True)
            doc.status = "failed"
            doc.error_msg = str(e)
            await db.commit()


def _extract_text(file_path: str, file_type: str) -> str:
    if file_type == "pdf":
        # Try multiple PDF libraries in order of preference
        # pypdf is lightest; pdfplumber handles complex layouts better
        text = ""

        # Method 1: pypdf (fast, handles most PDFs)
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if text.strip():
                logger.info(f"PDF extracted via pypdf: {len(text)} chars, {len(reader.pages)} pages")
                return text
        except ImportError:
            logger.debug("pypdf not installed, trying pdfplumber")
        except Exception as e:
            logger.warning(f"pypdf failed: {e}, trying pdfplumber")

        # Method 2: pdfplumber (better for tables/complex layouts)
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                pages_text = [p.extract_text() or "" for p in pdf.pages]
                text = "\n".join(pages_text)
            if text.strip():
                logger.info(f"PDF extracted via pdfplumber: {len(text)} chars")
                return text
        except ImportError:
            logger.debug("pdfplumber not installed")
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")

        # Method 3: PyMuPDF / fitz (handles scanned PDFs with OCR layer)
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            if text.strip():
                logger.info(f"PDF extracted via PyMuPDF: {len(text)} chars")
                return text
        except ImportError:
            logger.debug("PyMuPDF not installed")
        except Exception as e:
            logger.warning(f"PyMuPDF failed: {e}")

        if not text.strip():
            logger.error(f"All PDF extraction methods failed for {file_path}")
        return text

    elif file_type == "docx":
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""

    elif file_type in ("txt", "csv"):
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    return ""